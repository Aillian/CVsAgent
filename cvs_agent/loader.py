"""Document loader for PDF and DOCX CVs.

Supports:
- PDF (via pypdf)
- DOCX (via python-docx)
- Optional OCR fallback for scanned/image-only PDFs (via pytesseract + pdf2image)

Loading is parallelised with a thread pool because it is I/O-bound.
"""
from __future__ import annotations

import concurrent.futures
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from .config import MAX_PDF_BYTES, SUPPORTED_EXTENSIONS
from .console import console


@dataclass
class LoadedDocument:
    """A parsed CV document with minimal metadata."""

    filename: str
    path: Path
    text: str
    sha256: Optional[str] = None  # populated later if caching is enabled


# --- Low-level parsers ---------------------------------------------------


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _parse_docx(path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # include table cell text too — many résumés use tables heavily
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _ocr_pdf(path: Path) -> str:
    """Run OCR on a PDF file. Only called if text extraction yields empty output.

    Requires the optional dependencies `pytesseract` and `pdf2image`, plus the
    system `tesseract` and `poppler` binaries.
    """
    try:
        import pytesseract  # type: ignore
        from pdf2image import convert_from_path  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional path
        raise RuntimeError(
            "OCR requested but pytesseract/pdf2image are not installed. "
            "Install with: pip install pytesseract pdf2image"
        ) from exc

    images = convert_from_path(str(path))
    return "\n".join(pytesseract.image_to_string(img) for img in images).strip()


# --- Public API ----------------------------------------------------------


def load_single_document(path: Path, ocr_fallback: bool = False) -> Optional[LoadedDocument]:
    """Load a single PDF/DOCX file into a :class:`LoadedDocument`."""
    try:
        if not path.exists():
            console.warn("Loader", f"File not found: {path.name}")
            return None

        size = path.stat().st_size
        if size == 0:
            console.warn("Loader", f"Empty file skipped: {path.name}")
            return None
        if size > MAX_PDF_BYTES:
            console.warn(
                "Loader",
                f"File too large ({size / 1_048_576:.1f} MB, limit "
                f"{MAX_PDF_BYTES / 1_048_576:.0f} MB): {path.name}",
            )
            return None

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = _parse_pdf(path)
            if not text and ocr_fallback:
                console.log("Loader", f"No text in {path.name}; running OCR fallback", style="yellow")
                text = _ocr_pdf(path)
        elif suffix == ".docx":
            text = _parse_docx(path)
        else:
            console.warn("Loader", f"Unsupported extension: {path.name}")
            return None

        if not text:
            console.warn("Loader", f"No text extracted from {path.name}")
            return None

        return LoadedDocument(filename=path.name, path=path, text=text)
    except Exception as exc:  # noqa: BLE001 - surface any parser error
        console.error("Loader", f"Failed to load {path.name}: {exc}")
        return None


def discover_cv_files(directory: Path) -> List[Path]:
    """Return a sorted list of supported CV files in ``directory``."""
    if not directory.exists():
        console.error("Loader", f"Directory not found: {directory}")
        return []
    files = [
        p for p in directory.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return sorted(files, key=lambda p: p.name.lower())


def load_cvs(
    directory: Path,
    ocr_fallback: bool = False,
    max_workers: Optional[int] = None,
) -> List[LoadedDocument]:
    """Load every supported CV from ``directory`` in parallel."""
    files = discover_cv_files(directory)
    if not files:
        return []

    docs: List[LoadedDocument] = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        for doc in executor.map(lambda p: load_single_document(p, ocr_fallback), files):
            if doc is not None:
                docs.append(doc)
                console.log("Loader", f"Loaded: {doc.filename}")
    return docs
